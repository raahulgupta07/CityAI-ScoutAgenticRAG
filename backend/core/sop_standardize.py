"""
SOP Standardizer — Tier-1 Consulting Grade
Combines: McKinsey Pyramid Principle, Deloitte ITSM Delivery, Accenture ADM, PwC Controls
Generates DOCX with embedded screenshots, Mermaid workflows, RACI, control points.
Preserves original author + AI co-creator credit.
"""
from __future__ import annotations

import json, os, base64, urllib.request, logging, tempfile
from pathlib import Path
from io import BytesIO
from typing import Optional, Callable

from backend.core import database as db
from backend.core.config import call_openrouter, ROUTER_MODEL

logger = logging.getLogger(__name__)


def _get_template_context(department: str) -> str:
    """Get template-specific AI prompt section based on department."""
    try:
        from backend.core.sop_templates import get_template_for_department, get_template_prompt_section
        template_id = get_template_for_department(department)
        return get_template_prompt_section(template_id)
    except Exception:
        return ""


def _call_llm(prompt: str, max_tokens: int = 8000) -> str:
    return call_openrouter(prompt, model=ROUTER_MODEL, max_tokens=max_tokens)


# ── AI: Tier-1 Consulting Structure ──────────────────────────────────────────

def _prepare_page_content(pages: list) -> tuple[list[str], list[int]]:
    """Extract text content and image-only page numbers from page records."""
    content_parts, image_only_pages = [], []
    for p in pages:
        pg = p.get("page", 0)
        text = p.get("enhanced_content") or p.get("vision_content") or p.get("text_content") or ""
        src = "enhanced" if p.get("enhanced_content") else "vision" if p.get("vision_content") else "text"
        if src == "vision" and not p.get("text_content"):
            image_only_pages.append(pg)
        if text.strip():
            content_parts.append(f"[Page {pg} ({src})]\n{text.strip()[:2000]}")
    return content_parts, image_only_pages


def _parse_llm_json(result_text: str) -> dict:
    """Parse JSON from LLM response, stripping markdown fences."""
    result_text = result_text.strip()
    if result_text.startswith("```"):
        lines = result_text.split("\n")
        result_text = "\n".join(lines[1:])
        if result_text.rstrip().endswith("```"):
            result_text = result_text.rstrip()[:-3]
        result_text = result_text.strip()
    return json.loads(result_text)


# ── Chunked processing constants ──────────────────────────────────────────────
# Pages per chunk — tuned so each chunk fits within LLM context with prompt overhead
_PAGES_PER_CHUNK = 10
# Max chars per chunk of content sent to LLM
_CHUNK_CONTENT_LIMIT = 18000


def _build_continuation_prompt(title: str, department: str, chunk_content: str,
                                chunk_pages: str, step_offset: int) -> str:
    """Build prompt for continuation chunks — extract only procedures/definitions/escalation."""
    return f"""You are continuing to standardize a large document. Extract ONLY the procedural content from these pages.

DOCUMENT: {title}
DEPARTMENT: {department}
PAGES IN THIS SECTION: {chunk_pages}
CONTINUE STEP NUMBERING FROM: {step_offset + 1}

CONTENT:
{chunk_content}

CRITICAL LANGUAGE RULES:
- PRESERVE ALL LANGUAGES exactly as they appear
- If bilingual, keep BOTH languages

Extract procedures, definitions, escalation items, and references from THESE PAGES ONLY.
Return ONLY valid JSON (no markdown fences):
{{
  "procedure": [
    {{
      "step_number": {step_offset + 1},
      "title": "clear action title (verb-first)",
      "input": "what you need",
      "activity": "SHORT BULLET POINTS using '\\n- ' for each action",
      "output": "what is produced",
      "verification": "how to verify",
      "source_page": 0,
      "has_screenshot": false,
      "time_estimate": "duration",
      "control_point": false,
      "control_type": "Preventive|Detective|Corrective",
      "warnings": "",
      "notes": "",
      "decision_gate": false,
      "decision_options": []
    }}
  ],
  "definitions": [{{"term": "...", "definition": "..."}}],
  "escalation": [{{"trigger": "...", "action": "...", "escalated_to": "...", "timeframe": "...", "severity": "Critical|High|Medium|Low"}}],
  "references": ["any referenced documents found in these pages"],
  "kpis": [{{"metric": "...", "target": "...", "measurement": "...", "frequency": "..."}}],
  "raci": [{{"activity": "...", "responsible": "...", "accountable": "...", "consulted": "...", "informed": "..."}}],
  "ai_improvements": ["improvements for this section"]
}}"""


def _build_final_summary_prompt(title: str, department: str, total_pages: int,
                                 all_step_titles: list, all_definitions: list,
                                 image_only_pages: list, existing_summary: str) -> str:
    """Build prompt for final pass — regenerate executive summary and gap analysis over full doc."""
    steps_overview = "\n".join(f"  {i+1}. {t}" for i, t in enumerate(all_step_titles))
    defs_overview = ", ".join(d.get("term", "") for d in all_definitions[:30])
    return f"""You previously standardized the document below in chunks. Now write a final executive summary and gap analysis covering ALL {total_pages} pages.

DOCUMENT: {title}
DEPARTMENT: {department}
TOTAL PAGES: {total_pages}
IMAGE-ONLY PAGES: {image_only_pages}

ALL PROCEDURE STEPS EXTRACTED:
{steps_overview}

KEY DEFINITIONS: {defs_overview}

PREVIOUS SUMMARY (from first chunk only — may be incomplete):
{existing_summary}

Return ONLY valid JSON (no markdown fences):
{{
  "executive_summary": "McKinsey-style: 4-6 sentences covering the ENTIRE document. Lead with the conclusion — what this document achieves and why it matters. Mention the total scope ({total_pages} pages, {len(all_step_titles)} steps).",
  "mermaid_workflow": "flowchart TD\\n    A[Step] --> B{{{{Decision}}}}\\n    B -->|Yes| C[Action]\\n    B -->|No| D[Other]\\n    ... (cover ALL major steps)",
  "mermaid_swimlane": "flowchart TD\\n    subgraph Role1[Primary]\\n    A[First]\\n    end\\n    subgraph Role2[Secondary]\\n    B[Next]\\n    end\\n    A --> B\\n    ... (cover key roles)",
  "gap_analysis": {{
    "original_score": 0,
    "standardized_score": 0,
    "missing_sections": [],
    "image_only_pages": {image_only_pages},
    "improvements": [],
    "compliance_notes": ["ISO/ITIL compliance observations"]
  }}
}}"""


def analyze_and_structure(sop_id: str, tenant_id: str = None, on_status: Callable = None) -> dict:
    def _status(step, msg):
        if on_status: on_status(step, msg)

    sop = db.get_sop(sop_id, tenant_id=tenant_id)
    if not sop:
        return {"error": "Document not found"}
    pages = db.get_page_contents(sop_id, tenant_id=tenant_id)
    if not pages:
        return {"error": "No page content — process document first"}

    content_parts, image_only_pages = _prepare_page_content(pages)
    total_content_chars = sum(len(c) for c in content_parts)
    title = sop.get('title', sop_id)
    department = sop.get('department', 'Unknown')
    template_context = _get_template_context(department)

    # ── Single-pass for small documents (fits in one LLM call) ────────────
    if total_content_chars <= _CHUNK_CONTENT_LIMIT:
        _status("sop_standardize", f"Single-pass analysis ({len(pages)} pages, {total_content_chars:,} chars)")
        return _analyze_single_pass(sop, sop_id, tenant_id, pages, content_parts, image_only_pages, template_context)

    # ── Chunked processing for large documents ────────────────────────────
    num_chunks = (len(content_parts) + _PAGES_PER_CHUNK - 1) // _PAGES_PER_CHUNK
    _status("sop_standardize", f"Large document detected: {len(pages)} pages, {total_content_chars:,} chars")
    _status("sop_standardize", f"Processing in {num_chunks} chunks of ~{_PAGES_PER_CHUNK} pages each")

    chunks = [content_parts[i:i + _PAGES_PER_CHUNK] for i in range(0, len(content_parts), _PAGES_PER_CHUNK)]

    # ── Chunk 1: Full analysis (metadata + first batch of procedures) ─────
    _status("sop_standardize", f"Chunk 1/{num_chunks}: Full analysis (pages 1-{min(_PAGES_PER_CHUNK, len(pages))})")
    first_chunk_content = chr(10).join(chunks[0])[:_CHUNK_CONTENT_LIMIT]
    first_result = _analyze_single_pass(sop, sop_id, tenant_id, pages, chunks[0], image_only_pages, template_context)
    if "error" in first_result:
        return first_result

    merged = first_result
    step_offset = len(merged.get("procedure", []))

    # ── Chunks 2..N: Extract procedures/definitions/escalation ────────────
    for ci, chunk in enumerate(chunks[1:], start=2):
        chunk_content = chr(10).join(chunk)[:_CHUNK_CONTENT_LIMIT]
        # Extract page range from content headers like "[Page 11 (text)]"
        chunk_page_nums = []
        for part in chunk:
            import re as _re
            m = _re.match(r'\[Page (\d+)', part)
            if m: chunk_page_nums.append(m.group(1))
        page_range = f"{chunk_page_nums[0]}-{chunk_page_nums[-1]}" if chunk_page_nums else f"chunk {ci}"

        _status("sop_standardize", f"Chunk {ci}/{num_chunks}: Extracting procedures (pages {page_range})")

        prompt = _build_continuation_prompt(title, department, chunk_content, page_range, step_offset)
        try:
            result_text = _call_llm(prompt, max_tokens=8000)
            chunk_data = _parse_llm_json(result_text)
        except json.JSONDecodeError as e:
            logger.warning(f"Chunk {ci} JSON parse error: {e} — skipping")
            _status("sop_standardize", f"Chunk {ci}: parse error, skipping")
            continue
        except Exception as e:
            logger.warning(f"Chunk {ci} error: {e} — skipping")
            _status("sop_standardize", f"Chunk {ci}: error, skipping")
            continue

        # Merge procedures with correct step numbering
        new_steps = chunk_data.get("procedure", [])
        for step in new_steps:
            step_offset += 1
            step["step_number"] = step_offset
        merged.setdefault("procedure", []).extend(new_steps)

        # Merge lists (deduplicate definitions by term)
        existing_terms = {d.get("term", "").lower() for d in merged.get("definitions", [])}
        for d in chunk_data.get("definitions", []):
            if d.get("term", "").lower() not in existing_terms:
                merged.setdefault("definitions", []).append(d)
                existing_terms.add(d.get("term", "").lower())

        merged.setdefault("escalation", []).extend(chunk_data.get("escalation", []))
        merged.setdefault("references", []).extend(chunk_data.get("references", []))
        merged.setdefault("kpis", []).extend(chunk_data.get("kpis", []))
        merged.setdefault("raci", []).extend(chunk_data.get("raci", []))
        merged.setdefault("ai_improvements", []).extend(chunk_data.get("ai_improvements", []))

        _status("sop_standardize", f"Chunk {ci}: +{len(new_steps)} steps, total {step_offset} steps")

    # Deduplicate references
    if merged.get("references"):
        merged["references"] = list(dict.fromkeys(merged["references"]))

    # ── Final pass: Regenerate summary and diagrams over full content ──────
    _status("sop_standardize", f"Final pass: Regenerating executive summary and diagrams over all {step_offset} steps")
    all_step_titles = [s.get("title", f"Step {s.get('step_number', '?')}") for s in merged.get("procedure", [])]
    try:
        final_prompt = _build_final_summary_prompt(
            title, department, len(pages), all_step_titles,
            merged.get("definitions", []), image_only_pages,
            merged.get("executive_summary", "")
        )
        final_text = _call_llm(final_prompt, max_tokens=4000)
        final_data = _parse_llm_json(final_text)
        # Overwrite with full-document summary
        merged["executive_summary"] = final_data.get("executive_summary", merged.get("executive_summary", ""))
        merged["mermaid_workflow"] = final_data.get("mermaid_workflow", merged.get("mermaid_workflow", ""))
        merged["mermaid_swimlane"] = final_data.get("mermaid_swimlane", merged.get("mermaid_swimlane", ""))
        merged["gap_analysis"] = final_data.get("gap_analysis", merged.get("gap_analysis", {}))
    except Exception as e:
        logger.warning(f"Final summary pass failed: {e} — keeping chunk 1 summary")
        _status("sop_standardize", f"Final summary: using chunk 1 summary (final pass error: {e})")

    merged["total_pages"] = len(pages)
    merged["chunks_processed"] = num_chunks
    _status("sop_standardize", f"All {num_chunks} chunks merged: {step_offset} steps, {len(merged.get('definitions', []))} definitions")
    return merged


def _analyze_single_pass(sop: dict, sop_id: str, tenant_id: str, pages: list,
                          content_parts: list, image_only_pages: list, template_context: str) -> dict:
    """Original single-pass analysis for small documents or first chunk of large docs."""
    prompt = f"""You are a Tier-1 management consulting document standardization expert (McKinsey, Deloitte, Accenture, PwC combined).
Analyze this document and create a world-class standardized document.

DOCUMENT: {sop.get('title', sop_id)}
DEPARTMENT: {sop.get('department', 'Unknown')}
PAGES: {len(pages)} | IMAGE-ONLY PAGES: {image_only_pages}

{template_context}

CONTENT:
{chr(10).join(content_parts)[:_CHUNK_CONTENT_LIMIT]}

APPLY THESE FRAMEWORKS:
1. McKINSEY PYRAMID: Start every section with the conclusion/key insight FIRST, then supporting details
2. DELOITTE ITSM: Include RACI matrix, SLA targets, process KPIs
3. ACCENTURE ADM: Every step must have Input → Activity → Output → Verification
4. PwC CONTROLS: Mark control points and audit checkpoints in the procedure

CRITICAL LANGUAGE RULES:
- PRESERVE ALL LANGUAGES exactly as they appear in the original document
- If the document has Burmese/Myanmar text alongside English, KEEP BOTH languages in the output
- Do NOT translate or remove non-English text — include it exactly as-is
- For bilingual documents: put the English text first, then the original language text in parentheses or on the next line
- The procedure steps should contain BOTH languages if the original had both

EXTRACT from the actual content — preserve all original information including all languages. For image-only pages, write detailed step instructions from the vision descriptions.

Return ONLY valid JSON (no markdown fences):
{{
  "title": "document title",
  "subtitle": "one-line value proposition of this SOP",
  "executive_summary": "McKinsey-style: 3-4 sentences. Lead with the conclusion — what this SOP achieves and why it matters. Then the supporting logic.",
  "sop_number": "DOC-XXX-NNN",
  "version": "1.0",
  "department": "department",
  "category": "category",
  "original_author": "real author name from document (look for author/created by/prepared by)",
  "original_date": "date from document if found",
  "effective_date": "date if found",
  "classification": "Public|Internal|Confidential|Restricted",
  "review_cycle": "Quarterly|Semi-Annual|Annual",
  "purpose": "Extract the EXACT purpose from the original document. Quote or closely paraphrase the original text — do NOT write a generic purpose. If the original says 'This SOP provides troubleshooting steps for Line Manager Not Found error' then use THAT exact text.",
  "scope": {{
    "governing_thought": "McKinsey: one sentence that captures the entire scope",
    "in_scope": ["list"],
    "out_of_scope": ["list"],
    "applicable_to": ["roles/departments this applies to"]
  }},
  "kpis": [
    {{"metric": "metric name", "target": "target value", "measurement": "how measured", "frequency": "how often"}}
  ],
  "raci": [
    {{"activity": "activity name", "responsible": "who does it", "accountable": "who owns it", "consulted": "who to ask", "informed": "who to tell"}}
  ],
  "definitions": [{{"term": "...", "definition": "..."}}],
  "prerequisites": ["what must be in place before starting"],
  "procedure": [
    {{
      "step_number": 1,
      "title": "clear action title (verb-first)",
      "input": "what you need before this step",
      "activity": "Write as SHORT BULLET POINTS, not paragraphs. Use '\\n- ' for each action. Example: '- Open the System Settings\\n- Navigate to User Management\\n- Click on the affected user profile\\n- Check the Line Manager field'",
      "output": "what is produced by this step",
      "verification": "how to verify this step was done correctly",
      "source_page": 1,
      "has_screenshot": true,
      "time_estimate": "estimated duration",
      "control_point": true,
      "control_type": "Preventive|Detective|Corrective",
      "warnings": "cautions or risks",
      "notes": "additional context",
      "decision_gate": false,
      "decision_options": []
    }}
  ],
  "escalation": [{{"trigger": "...", "action": "...", "escalated_to": "...", "timeframe": "...", "severity": "Critical|High|Medium|Low"}}],
  "references": ["referenced documents"],
  "mermaid_workflow": "flowchart TD\\n    A[Step] --> B{{Decision}}\\n    B -->|Yes| C[Action]\\n    B -->|No| D[Other]\\n    ...",
  "mermaid_swimlane": "flowchart TD\\n    subgraph Role1[Service Desk]\\n    A[Log]\\n    end\\n    subgraph Role2[L2 Support]\\n    B[Diagnose]\\n    end\\n    A --> B\\n    ...",
  "ai_improvements": ["specific improvements made"],
  "gap_analysis": {{
    "original_score": 0,
    "standardized_score": 0,
    "missing_sections": [],
    "image_only_pages": {image_only_pages},
    "improvements": [],
    "compliance_notes": ["ISO/ITIL compliance observations"]
  }}
}}"""

    try:
        result_text = _call_llm(prompt, max_tokens=10000)
        structured = _parse_llm_json(result_text)
        structured["sop_id"] = sop_id
        structured["tenant_id"] = tenant_id
        structured["total_pages"] = len(pages)
        structured["image_only_count"] = len(image_only_pages)
        # Force real title and metadata from DB — don't let AI rename
        structured["title"] = sop.get("title") or structured.get("title", sop_id)
        structured["department"] = sop.get("department") or structured.get("department", "")
        structured["category"] = sop.get("category") or structured.get("category", "")
        structured["sop_number"] = sop_id  # Use actual document ID
        return structured
    except json.JSONDecodeError as e:
        logger.error(f"JSON parse error: {e}")
        return {"error": f"AI response not valid JSON: {str(e)[:200]}"}
    except Exception as e:
        logger.error(f"Structuring error: {e}")
        return {"error": str(e)}


# ── Mermaid ──────────────────────────────────────────────────────────────────

def render_mermaid_png(code: str) -> Optional[bytes]:
    try:
        enc = base64.urlsafe_b64encode(code.encode()).decode()
        req = urllib.request.Request(f"https://mermaid.ink/img/{enc}", headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as r:
            return r.read()
    except Exception as e:
        logger.warning(f"Mermaid render failed: {e}")
        return None


# ── DOCX Generation — Tier-1 Consulting Grade ────────────────────────────────

def generate_docx(structured: dict, sop_id: str, tenant_id: str = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.9)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.7)

    # Header/Footer will be set after we know the title
    def _setup_header_footer(sop_num, title, classification):
        from docx.oxml.ns import qn as _qn
        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hr = hp.add_run(f'{sop_num}  |  {classification}')
        hr.font.size = Pt(7.5); hr.font.color.rgb = RGBColor(120, 130, 140); hr.font.name = 'Calibri'
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        # Bottom border on header
        hpPr = hp._p.get_or_add_pPr()
        hpPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="B8CDE0"/></w:pBdr>'))

        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fr1 = fp.add_run(f'{title}')
        fr1.font.size = Pt(7); fr1.font.color.rgb = RGBColor(120, 130, 140); fr1.font.name = 'Calibri'
        fp.add_run('  |  ').font.size = Pt(7)
        # Page number field
        fp.add_run('Page ').font.size = Pt(7)
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        fp.add_run()._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        fp.add_run()._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        fp.add_run()._r.append(fldChar2)
        fp.add_run(' of ').font.size = Pt(7)
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        fp.add_run()._r.append(fldChar3)
        instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
        fp.add_run()._r.append(instrText2)
        fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        fp.add_run()._r.append(fldChar4)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        # Top border on footer
        fpPr = fp._p.get_or_add_pPr()
        fpPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="4" w:color="B8CDE0"/></w:pBdr>'))

    # Styles
    sn = doc.styles['Normal']
    sn.font.name = 'Calibri'; sn.font.size = Pt(10); sn.font.color.rgb = RGBColor(33, 33, 33)
    sn.paragraph_format.space_after = Pt(4); sn.paragraph_format.line_spacing = 1.2

    heading_cfg = {1: (17, (12, 50, 85)), 2: (13, (18, 70, 115)), 3: (11, (25, 90, 140))}
    for lvl, (sz, clr) in heading_cfg.items():
        h = doc.styles[f'Heading {lvl}']
        h.font.name = 'Calibri'; h.font.size = Pt(sz); h.font.bold = True
        h.font.color.rgb = RGBColor(*clr)
        h.paragraph_format.space_before = Pt(16 if lvl == 1 else 12)
        h.paragraph_format.space_after = Pt(6)

    # Colors
    NAVY, DARK, MID, LIGHT, BORDER = "0C3254", "1A4672", "2980B9", "EBF2F8", "B8CDE0"
    WARN_BG, WARN_BD = "FFF8E1", "F9A825"
    OK_BG, OK_BD = "E8F5E9", "2E7D32"
    CTRL_BG, CTRL_BD = "F3E5F5", "7B1FA2"

    # ── Helpers ──
    def P(text, bold=False, size=None, color=None, align=None, sb=None, sa=None, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        if bold: r.bold = True
        if italic: r.italic = True
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor(*color)
        r.font.name = 'Calibri'
        if align: p.alignment = align
        if sb is not None: p.paragraph_format.space_before = Pt(sb)
        if sa is not None: p.paragraph_format.space_after = Pt(sa)
        return p

    def bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        r = p.add_run(text); r.font.size = Pt(10); r.font.name = 'Calibri'

    def line():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="{BORDER}"/></w:pBdr>'))

    def tbl(headers, rows, header_color=NAVY):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = True
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = ''
            r = c.paragraphs[0].add_run(h)
            r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(255, 255, 255); r.font.name = 'Calibri'
            c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_color}"/>'))
        for j, rd in enumerate(rows):
            for i, v in enumerate(rd):
                c = t.rows[j+1].cells[i]; c.text = ''
                r = c.paragraphs[0].add_run(str(v)); r.font.size = Pt(8.5); r.font.name = 'Calibri'
                if j % 2 == 0:
                    c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT}"/>'))
        tblPr = t._tbl.tblPr or parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tblPr.append(parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="6" w:space="0" w:color="{BORDER}"/><w:left w:val="single" w:sz="6" w:space="0" w:color="{BORDER}"/><w:bottom w:val="single" w:sz="6" w:space="0" w:color="{BORDER}"/><w:right w:val="single" w:sz="6" w:space="0" w:color="{BORDER}"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER}"/><w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER}"/></w:tblBorders>'))
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def callout(text, bg, border):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="20" w:space="8" w:color="{border}"/></w:pBdr>'))
        pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>'))
        p.paragraph_format.left_indent = Inches(0.12)
        r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = RGBColor(33, 33, 33); r.font.name = 'Calibri'

    def _fit_image(img_path, max_page_w=6.5):
        """Calculate image size: keep original, only shrink if wider than page."""
        from PIL import Image as PILImage
        try:
            img = PILImage.open(img_path)
            w_px, h_px = img.size
            dpi = img.info.get('dpi', (96, 96))
            dpi_x = dpi[0] if isinstance(dpi, tuple) else 96
            if dpi_x < 10: dpi_x = 96  # sanity check
            w_in = w_px / dpi_x
            h_in = h_px / dpi_x
            if w_in > max_page_w:
                scale = max_page_w / w_in
                w_in = max_page_w
                h_in = h_in * scale
            return Inches(w_in), Inches(h_in)
        except Exception:
            return Inches(5.0), None  # fallback

    def diagram(png, caption):
        if not png: return
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False); tmp.write(png); tmp.close()
        w, h = _fit_image(tmp.name, max_page_w=6.0)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if h:
            p.add_run().add_picture(tmp.name, width=w, height=h)
        else:
            p.add_run().add_picture(tmp.name, width=w)
        os.unlink(tmp.name)
        P(caption, italic=True, size=8.5, color=(90, 100, 110), align=WD_ALIGN_PARAGRAPH.CENTER, sa=10)

    # Extract values
    title = structured.get("title", sop_id)
    subtitle = structured.get("subtitle", "")
    sop_num = structured.get("sop_number", f"SOP-{sop_id[:8].upper()}")
    version = structured.get("version", "1.0")
    dept = structured.get("department", "")
    author = structured.get("original_author", "Original Author")
    orig_date = structured.get("original_date", "")
    eff_date = structured.get("effective_date", "")
    classification = structured.get("classification", "Internal")
    review_cycle = structured.get("review_cycle", "Annual")
    agent_name = "Scout Agent"
    if tenant_id:
        try:
            t = db.get_tenant(tenant_id)
            if t: agent_name = t.get("agent_name", "Scout Agent")
        except Exception as _e: logger.debug(f"Non-critical: {_e}")

    # Setup header/footer on every page
    _setup_header_footer(sop_num, title, classification)

    # ══════════════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    for _ in range(3): P('', sa=0)

    # Top band
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p._p.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{NAVY}"/>'))
    r = p.add_run(f'   {classification.upper()}   '); r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(255, 255, 255); r.font.name = 'Calibri'

    P('', sa=8)
    P('STANDARD OPERATING PROCEDURE', bold=True, size=10, color=(12, 50, 85), align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    P(title, bold=True, size=24, color=(12, 50, 85), align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    if subtitle:
        P(subtitle, italic=True, size=11, color=(90, 100, 110), align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
    line()

    # Metadata grid (2-col table)
    mt = doc.add_table(rows=6, cols=4); mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    upload_dt = structured.get("upload_date", "—")
    trained_dt = structured.get("trained_date", "—")
    standardized_dt = structured.get("standardized_date", "—")
    meta = [
        ("Document ID", sop_num, "Department", dept),
        ("Version", version, "Classification", classification),
        ("Author", author, "Review Cycle", review_cycle),
        ("Co-Author", f"{agent_name} (AI)", "Total Pages", str(structured.get("total_pages", "—"))),
        ("Uploaded", upload_dt, "Trained", trained_dt),
        ("Standardized", standardized_dt, "Effective Date", eff_date or "—"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(meta):
        for ci, (lbl, val) in enumerate([(l1, v1), (l2, v2)]):
            lc = mt.rows[i].cells[ci * 2]; lc.text = ''
            r = lc.paragraphs[0].add_run(lbl); r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(80, 90, 100); r.font.name = 'Calibri'
            lc._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT}"/>'))
            vc = mt.rows[i].cells[ci * 2 + 1]; vc.text = ''
            r = vc.paragraphs[0].add_run(val); r.font.size = Pt(8); r.font.name = 'Calibri'

    P('', sa=6)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  EXECUTIVE SUMMARY (McKinsey Pyramid)
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('Executive Summary', level=1)
    callout(structured.get("executive_summary", "Executive summary not available."), LIGHT, DARK)
    P('', sa=4)

    # ══════════════════════════════════════════════════════════════════════════
    #  1. PURPOSE & SCOPE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('1. Purpose & Scope', level=1)

    doc.add_heading('1.1 Purpose', level=2)
    P(structured.get("purpose", "—"), size=10, sa=8)

    scope = structured.get("scope", {})
    gt = scope.get("governing_thought", "")
    if gt:
        doc.add_heading('1.2 Scope — Governing Thought', level=2)
        callout(gt, LIGHT, MID)

    ins = scope.get("in_scope", [])
    outs = scope.get("out_of_scope", [])
    if ins or outs:
        rows = []
        def _s(v):
            return str(v) if not isinstance(v, dict) else v.get("description", v.get("name", str(v)))
        for i in range(max(len(ins), len(outs), 1)):
            rows.append([_s(ins[i]) if i < len(ins) else "", _s(outs[i]) if i < len(outs) else ""])
        tbl(['In Scope', 'Out of Scope'], rows)

    applicable = scope.get("applicable_to", [])
    if applicable:
        P('Applicable To:', bold=True, size=9, color=(12, 50, 85), sa=2)
        for a in (applicable if isinstance(applicable, list) else [applicable]):
            bullet(str(a) if not isinstance(a, dict) else a.get("name", a.get("role", str(a))))

    # Workflow
    mermaid_wf = structured.get("mermaid_workflow", "")
    if mermaid_wf:
        doc.add_heading('1.3 Process Overview', level=2)
        diagram(render_mermaid_png(mermaid_wf), f'Figure 1: {title} — End-to-End Process Flow')

    # Swimlane diagram removed — one workflow diagram is sufficient

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  2. KPIs & SLA TARGETS (Deloitte)
    # ══════════════════════════════════════════════════════════════════════════
    kpis = structured.get("kpis", [])
    if kpis:
        doc.add_heading('2. Key Performance Indicators', level=1)
        P('Deloitte ITSM Delivery Framework — measurable targets tied to each process area.', italic=True, size=9, color=(90, 100, 110), sa=8)
        tbl(['Metric', 'Target', 'Measurement Method', 'Frequency'],
            [[k.get("metric", ""), k.get("target", ""), k.get("measurement", ""), k.get("frequency", "")] for k in kpis])

    # ══════════════════════════════════════════════════════════════════════════
    #  3. RACI MATRIX (Deloitte)
    # ══════════════════════════════════════════════════════════════════════════
    raci = structured.get("raci", [])
    if raci:
        doc.add_heading('3. RACI Matrix', level=1)
        P('R = Responsible (does the work)  |  A = Accountable (owns the outcome)  |  C = Consulted  |  I = Informed',
          italic=True, size=8.5, color=(90, 100, 110), sa=8)
        tbl(['Activity', 'Responsible', 'Accountable', 'Consulted', 'Informed'],
            [[r.get("activity", ""), r.get("responsible", ""), r.get("accountable", ""), r.get("consulted", ""), r.get("informed", "")] for r in raci],
            header_color=DARK)

    # ══════════════════════════════════════════════════════════════════════════
    #  4. DEFINITIONS
    # ══════════════════════════════════════════════════════════════════════════
    defs = structured.get("definitions", [])
    if defs:
        doc.add_heading('4. Definitions & Terminology', level=1)
        tbl(['Term', 'Definition'], [[d.get("term", ""), d.get("definition", "")] for d in defs])

    # ══════════════════════════════════════════════════════════════════════════
    #  5. PREREQUISITES
    # ══════════════════════════════════════════════════════════════════════════
    prereqs = structured.get("prerequisites", [])
    if prereqs:
        doc.add_heading('5. Prerequisites', level=1)
        for pr in prereqs:
            bullet(str(pr) if not isinstance(pr, dict) else pr.get("description", pr.get("name", str(pr))))
        P('', sa=4)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  6. PROCEDURE (Accenture Input→Activity→Output→Verification + PwC Controls)
    # ══════════════════════════════════════════════════════════════════════════
    procedure = structured.get("procedure", [])
    if procedure:
        doc.add_heading('6. Procedure', level=1)
        P(f'{len(procedure)} steps | Accenture ADM: Input → Activity → Output → Verification | PwC Control Points marked',
          italic=True, size=9, color=(90, 100, 110), sa=10)

        # Get screenshots — db returns dict keyed by page str, resolve full paths
        screenshots = {}
        try:
            ss_dict = db.get_screenshots(sop_id, tenant_id=tenant_id)
            # Try tenant dir first, fallback to public dir
            ss_base_tenant = db.DATA_DIR / "tenants" / tenant_id / "screenshots" / sop_id if tenant_id else None
            ss_base_public = db.DATA_DIR / "screenshots" / sop_id
            for page_str, ss_list in ss_dict.items():
                pg = int(page_str)
                if ss_list:
                    rel_path = ss_list[0].get("path", "")
                    # Try tenant path first, then public
                    full_path = None
                    if ss_base_tenant and (ss_base_tenant / rel_path).exists():
                        full_path = str(ss_base_tenant / rel_path)
                    elif (ss_base_public / rel_path).exists():
                        full_path = str(ss_base_public / rel_path)
                    if full_path:
                        screenshots[pg] = full_path
        except Exception as e:
            logger.warning(f"Screenshot loading: {e}")

        for step in procedure:
            snum = step.get("step_number", 0)
            stitle = step.get("title", "")
            activity = step.get("activity", step.get("description", ""))
            s_input = step.get("input", "")
            s_output = step.get("output", "")
            verification = step.get("verification", "")
            source_page = step.get("source_page", 0)
            has_ss = step.get("has_screenshot", False)
            time_est = step.get("time_estimate", "")
            is_control = step.get("control_point", False)
            ctrl_type = step.get("control_type", "")
            warnings = step.get("warnings", "")
            notes = step.get("notes", "")
            is_decision = step.get("decision_gate", False)
            dec_options = step.get("decision_options", [])

            # Step header with optional control point badge
            h = doc.add_heading(f'6.{snum}  {stitle}', level=2)
            if time_est:
                P(f'Estimated Duration: {time_est}', italic=True, size=8.5, color=(90, 100, 110), sa=2)

            # Control point marker
            if is_control:
                callout(f'CONTROL POINT ({ctrl_type}): This step requires verification before proceeding.', CTRL_BG, CTRL_BD)

            # Accenture ADM table: Input → Activity → Output → Verification
            adm = doc.add_table(rows=4, cols=2); adm.alignment = WD_TABLE_ALIGNMENT.CENTER; adm.autofit = True
            adm_data = [("INPUT", s_input or "Previous step output"), ("ACTIVITY", activity),
                        ("OUTPUT", s_output or "Completed action"), ("VERIFICATION", verification or "Visual confirmation")]
            for i, (lbl, val) in enumerate(adm_data):
                lc = adm.rows[i].cells[0]; lc.text = ''
                r = lc.paragraphs[0].add_run(lbl); r.bold = True; r.font.size = Pt(8); r.font.name = 'Calibri'
                r.font.color.rgb = RGBColor(255, 255, 255)
                colors = [NAVY, DARK, MID, "2E7D32"]
                lc._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors[i]}"/>'))
                vc = adm.rows[i].cells[1]; vc.text = ''
                # Handle bullet points in activity — split by newline and '- '
                val_str = str(val)
                lines = [l.strip().lstrip('- ').strip() for l in val_str.split('\n') if l.strip()]
                if len(lines) > 1:
                    # Multiple lines → render as bullet points
                    first = True
                    for line_text in lines:
                        if not line_text: continue
                        if first:
                            r = vc.paragraphs[0].add_run(f'• {line_text}')
                            r.font.size = Pt(9); r.font.name = 'Calibri'
                            first = False
                        else:
                            bp = vc.add_paragraph()
                            bp.paragraph_format.space_before = Pt(1)
                            bp.paragraph_format.space_after = Pt(1)
                            r = bp.add_run(f'• {line_text}')
                            r.font.size = Pt(9); r.font.name = 'Calibri'
                else:
                    r = vc.paragraphs[0].add_run(val_str); r.font.size = Pt(9); r.font.name = 'Calibri'
                if i % 2 == 0:
                    vc._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT}"/>'))
            P('', sa=4)

            # Screenshot
            if has_ss and source_page in screenshots:
                ss_path = screenshots[source_page]
                if os.path.exists(ss_path):
                    try:
                        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Keep original size, only cap at page width
                        w, h = _fit_image(ss_path, max_page_w=6.0)
                        if h:
                            p.add_run().add_picture(ss_path, width=w, height=h)
                        else:
                            p.add_run().add_picture(ss_path, width=w)
                        P(f'Figure {snum}: {stitle} (Source: Page {source_page})',
                          italic=True, size=8, color=(90, 100, 110), align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
                    except Exception as _e: logger.debug(f"Non-critical: {_e}")

            # Decision gate
            if is_decision and dec_options:
                callout('DECISION GATE: ' + ' | '.join(str(o) if not isinstance(o, dict) else o.get("option", str(o)) for o in dec_options), WARN_BG, WARN_BD)

            # Warnings
            if warnings:
                callout(f'WARNING: {warnings}', WARN_BG, WARN_BD)

            # Notes
            if notes:
                P(f'Note: {notes}', italic=True, size=9, color=(90, 100, 110), sa=6)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  7. ESCALATION
    # ══════════════════════════════════════════════════════════════════════════
    escalation = structured.get("escalation", [])
    if escalation:
        doc.add_heading('7. Escalation Matrix', level=1)
        tbl(['Trigger', 'Severity', 'Action', 'Escalated To', 'Timeframe'],
            [[e.get("trigger", ""), e.get("severity", ""), e.get("action", ""), e.get("escalated_to", ""), e.get("timeframe", "")]
             for e in escalation])

    # ══════════════════════════════════════════════════════════════════════════
    #  REFERENCES
    # ══════════════════════════════════════════════════════════════════════════
    refs = structured.get("references", [])
    if refs:
        doc.add_heading('References', level=1)
        for ref in refs: bullet(str(ref))
        P('', sa=4)

    # ══════════════════════════════════════════════════════════════════════════
    #  REVISION HISTORY
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('Document Control — Revision History', level=1)
    rev_rows = [
        [version, orig_date or '—', author, 'Original document created'],
        ['—', upload_dt, 'System', 'Document uploaded to Scout Agentic RAG platform'],
        ['—', trained_dt, 'AI Agent', 'Document processed, indexed, and agent trained'],
        [f'{version}-S', standardized_dt, f'{agent_name} (AI Co-Author)',
         'AI-standardized: Tier-1 consulting format, workflow diagrams, RACI matrix, bullet-point procedures, gap analysis'],
    ]
    tbl(['Version', 'Date', 'Author', 'Description of Changes'], rev_rows)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  AI ENHANCEMENT REPORT
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('AI Co-Author Enhancement Report', level=1)

    P(f'Original Author: {author}', bold=True, size=10, sa=2)
    P(f'AI Co-Author: {agent_name}', bold=True, size=10, color=(12, 50, 85), sa=2)
    P('Methodology: McKinsey Pyramid Principle | Deloitte ITSM Framework | Accenture ADM | PwC Controls',
      italic=True, size=9, color=(90, 100, 110), sa=12)

    gap = structured.get("gap_analysis", {})
    orig_score = gap.get("original_score", 0)
    std_score = gap.get("standardized_score", 0)

    # Score cards
    sc = doc.add_table(rows=2, cols=2); sc.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (label, score, bg, clr) in enumerate([
        ("ORIGINAL DOCUMENT", f"{orig_score}/100", "FFF3E0", (180, 80, 40)),
        ("STANDARDIZED (AI-ENHANCED)", f"{std_score}/100", "E8F5E9", (27, 122, 61))
    ]):
        hc = sc.rows[0].cells[ci]; hc.text = ''
        r = hc.paragraphs[0].add_run(label); r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(80, 80, 80); r.font.name = 'Calibri'
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hc._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>'))
        vc = sc.rows[1].cells[ci]; vc.text = ''
        r = vc.paragraphs[0].add_run(score); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(*clr); r.font.name = 'Calibri'
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vc._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>'))

    P('', sa=12)

    # Improvements
    improvements = structured.get("ai_improvements", [])
    if improvements:
        doc.add_heading('Improvements Made', level=2)
        for imp in improvements: bullet(str(imp))

    missing = gap.get("missing_sections", [])
    if missing:
        doc.add_heading('Sections Added by AI', level=2)
        for m in missing: callout(f'Added: {str(m) if not isinstance(m, dict) else m.get("section", str(m))}', WARN_BG, WARN_BD)

    img_pages = gap.get("image_only_pages", [])
    if img_pages:
        doc.add_heading('Image-Only Pages Enhanced', level=2)
        P(f'{len(img_pages)} pages contained only screenshots. AI generated detailed step descriptions using vision analysis.',
          size=10, sa=6)
        P(f'Pages: {", ".join(str(p) for p in img_pages)}', size=9, color=(90, 100, 110))

    comp = gap.get("compliance_notes", [])
    if comp:
        doc.add_heading('Compliance Observations', level=2)
        for c in comp: callout(str(c) if not isinstance(c, dict) else c.get("note", str(c)), LIGHT, MID)

    line()
    P(f'All original content by {author} has been preserved. The AI co-author restructured the format, added missing sections, '
      f'and applied Tier-1 consulting frameworks. No original procedures were altered — only enhanced with structure, '
      f'visual aids, and documentation standards.',
      size=8, color=(120, 130, 140), sa=0)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Full Pipeline ────────────────────────────────────────────────────────────

def standardize_sop(sop_id: str, tenant_id: str = None, on_status: Callable = None) -> dict:
    def _status(step, msg, *args):
        if on_status: on_status(step, msg)

    _status("sop_standardize", f"━━━ DOCUMENT STANDARDIZATION START ━━━")
    _status("sop_standardize", f"Document: {sop_id}")
    _status("sop_standardize", f"Frameworks: McKinsey Pyramid | Deloitte RACI | Accenture ADM | PwC Controls")
    _status("sop_standardize", f"Analyzing content and structuring via AI...")

    structured = analyze_and_structure(sop_id, tenant_id=tenant_id, on_status=on_status)
    if "error" in structured:
        _status("sop_error", f"Error: {structured['error']}")
        return structured

    # Add dates from DB
    sop = db.get_sop(sop_id, tenant_id=tenant_id)
    if sop:
        from datetime import datetime
        def _fmt_date(val):
            if not val: return "—"
            if isinstance(val, str):
                try: val = datetime.fromisoformat(val.replace("Z", "+00:00"))
                except: return val[:10] if len(val) >= 10 else val
            return val.strftime("%Y-%m-%d %H:%M") if hasattr(val, 'strftime') else str(val)[:16]
        structured["upload_date"] = _fmt_date(sop.get("created_at"))
        structured["trained_date"] = _fmt_date(sop.get("indexed_at"))
        structured["standardized_date"] = _fmt_date(datetime.now())

    author = structured.get("original_author", "Unknown")
    steps = len(structured.get("procedure", []))
    kpis = len(structured.get("kpis", []))
    raci = len(structured.get("raci", []))
    defs = len(structured.get("definitions", []))
    gap = structured.get("gap_analysis", {})
    orig_score = gap.get("original_score", 0)
    std_score = gap.get("standardized_score", 0)
    img_pages = len(gap.get("image_only_pages", []))

    _status("sop_standardize", f"AI Analysis Complete:")
    _status("sop_standardize", f"  Author detected: {author}")
    _status("sop_standardize", f"  Procedure steps: {steps}")
    _status("sop_standardize", f"  KPI metrics: {kpis}")
    _status("sop_standardize", f"  RACI activities: {raci}")
    _status("sop_standardize", f"  Definitions: {defs}")
    _status("sop_standardize", f"  Image-only pages enhanced: {img_pages}")
    _status("sop_standardize", f"  Score: {orig_score} → {std_score} /100")
    _status("sop_standardize", f"Rendering Mermaid workflow diagrams...")
    _status("sop_standardize", f"Generating Tier-1 DOCX with embedded screenshots...")

    try:
        docx_bytes = generate_docx(structured, sop_id, tenant_id=tenant_id)
    except Exception as e:
        logger.error(f"DOCX generation error: {e}")
        _status("sop_error", f"DOCX error: {e}")
        return {"error": f"DOCX generation failed: {e}"}

    data_dir = db.DATA_DIR / "tenants" / tenant_id / "standardized" if tenant_id else db.DATA_DIR / "standardized"
    data_dir.mkdir(parents=True, exist_ok=True)
    docx_path = data_dir / f"{sop_id}_standardized.docx"
    docx_path.write_bytes(docx_bytes)

    gap = structured.get("gap_analysis", {})
    sop_score = gap.get("standardized_score", 0)
    try:
        conn = db.get_db(tenant_id)
        conn.execute("UPDATE sops SET standardized_json = %s, sop_score = %s, sop_gaps = %s, standardized_at = NOW() WHERE sop_id = %s",
                     (json.dumps(structured), sop_score, json.dumps(gap), sop_id))
        conn.commit(); conn.close()
    except Exception as e:
        logger.warning(f"DB update failed: {e}")

    # Get agent name for status
    _agent_name = "Scout Agent"
    if tenant_id:
        try:
            _t = db.get_tenant(tenant_id)
            if _t: _agent_name = _t.get("agent_name", "Scout Agent")
        except Exception as _e: logger.debug(f"Non-critical: {_e}")

    _status("sop_done", f"━━━ DOCUMENT STANDARDIZATION COMPLETE ━━━")
    _status("sop_done", f"Score: {orig_score} → {sop_score}/100")
    _status("sop_done", f"Steps: {steps} | KPIs: {kpis} | RACI: {raci}")
    _status("sop_done", f"Author: {author}")
    _status("sop_done", f"Co-Author: {_agent_name} (AI)")
    _status("sop_done", f"DOCX: {docx_path.name} ({len(docx_bytes)//1024}KB)")
    _status("sop_done", f"Download ready!")

    return {"status": "standardized", "sop_id": sop_id, "score": sop_score,
            "steps": steps, "original_author": author, "gaps": gap, "docx_path": str(docx_path)}
