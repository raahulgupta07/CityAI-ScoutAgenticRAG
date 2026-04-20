"""
DOCX text + image extractor using python-docx.
No LibreOffice needed. Extracts directly from the native format.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Optional, Callable

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

from backend.core import database as db
from backend.core.database import SCREENSHOT_DIR, get_tenant_screenshot_dir


def extract_docx(file_path: str, sop_id: str, on_status: Optional[Callable] = None, tenant_id: str = None) -> dict:
    """
    Extract text + images from a DOCX file.
    Stores text in page_content table (1 section = 1 "page").
    Stores images in screenshots folder + table.
    """
    if on_status:
        on_status("docx", f"Extracting DOCX: {sop_id}")

    try:
        doc = Document(file_path)
    except Exception as e:
        return {"error": f"Cannot open DOCX: {e}"}

    # Extract all paragraphs grouped by headings (simulate "pages")
    pages = []
    current_page = {"page": 1, "text": "", "is_heading": False}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # New heading = new "page"
        if para.style.name.startswith("Heading"):
            if current_page["text"]:
                pages.append(current_page)
            current_page = {
                "page": len(pages) + 1,
                "text": f"**{text}**\n",
                "is_heading": True,
            }
        else:
            current_page["text"] += text + "\n"

    # Don't forget the last page
    if current_page["text"]:
        pages.append(current_page)

    # If no headings found, split by fixed line count
    if len(pages) <= 1 and current_page["text"]:
        lines = current_page["text"].split("\n")
        pages = []
        chunk_size = 40  # ~40 lines per "page"
        for i in range(0, len(lines), chunk_size):
            chunk = "\n".join(lines[i:i + chunk_size]).strip()
            if chunk:
                pages.append({"page": len(pages) + 1, "text": chunk})

    # Extract tables
    table_data = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
        rows = []
        for row in table.rows[1:]:
            rows.append([cell.text.strip() for cell in row.cells])
        table_data.append({"headers": headers, "rows": rows})

    # Store pages in page_content table
    for page in pages:
        page_num = page["page"]
        # Attach tables to first page
        page_tables = table_data if page_num == 1 else []

        db.upsert_page_content(
            sop_id=sop_id,
            page=page_num,
            text_content=page["text"],
            tables=page_tables,
            has_tables=len(page_tables) > 0,
            extraction_method="docx",
            tenant_id=tenant_id,
        )

    if on_status:
        on_status("docx_text", f"Extracted {len(pages)} sections from DOCX")

    # Extract embedded images — tenant-scoped directory
    screenshot_dir = get_tenant_screenshot_dir(tenant_id) / sop_id
    screenshot_dir.mkdir(parents=True, exist_ok=True)
    img_count = 0

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                img = Image.open(io.BytesIO(img_data))
                if img.width < 150 or img.height < 80:
                    continue

                img_count += 1
                if img.mode in ("RGBA", "P", "LA"):
                    img = img.convert("RGB")

                filename = f"docx_img{img_count}.png"
                save_path = screenshot_dir / filename
                img.save(str(save_path), "PNG")

                db.upsert_screenshot(sop_id, 1, img_count, filename, img.width, img.height, tenant_id=tenant_id)
            except Exception:
                continue

    # Generate HTML preview using mammoth
    try:
        import mammoth
        preview_dir = db.DATA_DIR / "previews"
        preview_dir.mkdir(parents=True, exist_ok=True)
        with open(file_path, "rb") as f:
            result = mammoth.convert_to_html(f)
            html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
            <style>body{{font-family:Inter,sans-serif;padding:40px;max-width:900px;margin:0 auto;color:#1a1a1a;line-height:1.7}}
            table{{border-collapse:collapse;width:100%;margin:16px 0}}th,td{{border:1px solid #e0e0e0;padding:8px 12px;text-align:left}}
            th{{background:#f5f5f5;font-weight:700}}img{{max-width:100%;border-radius:8px;margin:12px 0}}
            h1{{font-size:24px}}h2{{font-size:20px}}h3{{font-size:16px}}</style></head><body>{result.value}</body></html>"""
            (preview_dir / f"{sop_id}.html").write_text(html, encoding="utf-8")
        if on_status:
            on_status("docx_preview", f"HTML preview generated for {sop_id}")
    except Exception as e:
        if on_status:
            on_status("docx_preview_error", f"Preview generation failed: {e}")

    if on_status:
        on_status("docx_done", f"DOCX done: {len(pages)} sections, {len(table_data)} tables, {img_count} images")

    return {
        "pages": len(pages),
        "tables": len(table_data),
        "images": img_count,
    }
