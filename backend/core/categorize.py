"""
AI-powered document categorization.
Reads text from PDF/DOCX/XLSX and determines category, department, system, type, tags.
"""
from __future__ import annotations

import json
from pathlib import Path

from backend.core.config import get_openrouter_client, ROUTER_MODEL, INSTANCE

# Build categorize prompt from instance config
_categories = INSTANCE.get("categories", ["uncategorized"])
_doc_types = INSTANCE.get("document_types", ["guide", "policy", "report"])
_cat_examples = ", ".join(f'"{c}"' for c in _categories[:8])
_type_list = ", ".join(f'"{t}"' for t in _doc_types)

CATEGORIZE_PROMPT = f"""You are a document classifier. Given the first portion of a document, determine:

1. category: a path like {_cat_examples}
2. title: the document title/name
3. department: the department (e.g., "Group IT Technology", "Digital Product", "Energy", "Legal", "HR", "Finance")
4. system: the specific system if any (e.g., "Odoo", "GOLD", "City Family", "SAP", "Power BI") or ""
5. type: one of {_type_list}
6. tags: 3-5 relevant keywords

Return JSON only:
{{"category": "enterprise/odoo", "title": "POS Shop Setting for CFC", "department": "Group IT Technology", "system": "Odoo", "type": "configuration", "tags": ["POS", "shop", "CFC", "settings"]}}"""


def _extract_text(file_path: str) -> str:
    """Extract text from any supported file type."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        import PyPDF2
        try:
            reader = PyPDF2.PdfReader(file_path)
            text = ""
            for i in range(min(2, len(reader.pages))):
                text += reader.pages[i].extract_text() or ""
            return text
        except Exception:
            return ""

    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = []
            for p in doc.paragraphs[:60]:
                if p.text.strip():
                    paragraphs.append(p.text.strip())
            # Also grab table content
            for table in doc.tables[:5]:
                for row in table.rows[:10]:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            return "\n".join(paragraphs)
        except Exception:
            return ""

    elif ext in (".xlsx", ".xls"):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []
            for sheet_name in wb.sheetnames[:3]:
                ws = wb[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                for row in ws.iter_rows(max_row=30, values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        text_parts.append(" | ".join(cells))
            wb.close()
            return "\n".join(text_parts)
        except Exception:
            return ""

    return ""


def categorize_document(file_path: str) -> dict:
    """Read document content and categorize with AI. Works for PDF, DOCX, XLSX."""
    text = _extract_text(file_path)

    if not text.strip():
        return {
            "category": "uncategorized",
            "title": Path(file_path).stem,
            "department": "",
            "system": "",
            "type": "",
            "tags": [],
        }

    # Call LLM
    try:
        client = get_openrouter_client()
        response = client.chat.completions.create(
            model=ROUTER_MODEL,
            messages=[
                {"role": "system", "content": CATEGORIZE_PROMPT},
                {"role": "user", "content": f"Document text:\n{text[:3000]}"},
            ],
            max_tokens=300,
            temperature=0,
        )
        raw = response.choices[0].message.content.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        result = json.loads(raw.strip())
        return result
    except Exception as e:
        return {
            "category": "uncategorized",
            "title": Path(file_path).stem,
            "department": "",
            "system": "",
            "type": "",
            "tags": [],
            "error": str(e),
        }


# Backward compatibility alias
categorize_pdf = categorize_document
