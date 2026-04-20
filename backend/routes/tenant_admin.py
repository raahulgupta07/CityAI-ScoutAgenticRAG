"""Tenant-scoped admin API routes. All operations scoped to tenant's schema."""
from __future__ import annotations

import asyncio
import json
import re
import shutil
import threading
from queue import Queue
from typing import Optional
from pathlib import Path
from fastapi import APIRouter, Body, Query, UploadFile, File
from fastapi.responses import Response, StreamingResponse

from backend.core import database as db

router = APIRouter(prefix="/api/t/{tenant_id}/admin")

# Prevent concurrent processing of the same document
_processing_docs: set[str] = set()
_processing_lock = threading.Lock()


def _tenant_data_dir(tenant_id: str) -> Path:
    return db.DATA_DIR / "tenants" / tenant_id


@router.get("/search")
async def search_content(tenant_id: str, q: str = Query(""), limit: int = Query(20)):
    """Full-text search across all document page content and SOP metadata."""
    q = (q or "").strip()
    if len(q) < 2:
        return []
    limit = max(1, min(limit, 100))
    conn = db.get_db(tenant_id)
    try:
        cur = conn.cursor()
        # Search page_content using PostgreSQL full-text search
        cur.execute("""
            SELECT pc.sop_id, pc.page,
                   ts_headline('english', COALESCE(pc.enhanced_content, pc.vision_content, pc.text_content, ''),
                               plainto_tsquery('english', %s),
                               'StartSel=<mark>, StopSel=</mark>, MaxWords=50, MinWords=20') as snippet,
                   ts_rank(to_tsvector('english', COALESCE(pc.enhanced_content, pc.vision_content, pc.text_content, '')),
                           plainto_tsquery('english', %s)) as rank
            FROM page_content pc
            WHERE to_tsvector('english', COALESCE(pc.enhanced_content, pc.vision_content, pc.text_content, ''))
                  @@ plainto_tsquery('english', %s)
            ORDER BY rank DESC LIMIT %s
        """, (q, q, q, limit))
        content_rows = cur.fetchall()

        # Search sops table (title, sop_id) for matches
        cur.execute("""
            SELECT sop_id, title FROM sops
            WHERE to_tsvector('english', COALESCE(title, '') || ' ' || COALESCE(sop_id, ''))
                  @@ plainto_tsquery('english', %s)
               OR LOWER(title) LIKE %s
               OR LOWER(sop_id) LIKE %s
            LIMIT %s
        """, (q, f"%{q.lower()}%", f"%{q.lower()}%", limit))
        sop_rows = cur.fetchall()

        # Build title lookup from sop matches
        title_map = {r[0]: r[1] for r in sop_rows}

        # Also fetch titles for content results not already in the map
        missing_ids = [r[0] for r in content_rows if r[0] not in title_map]
        if missing_ids:
            cur.execute(
                "SELECT sop_id, title FROM sops WHERE sop_id = ANY(%s)",
                (missing_ids,)
            )
            for row in cur.fetchall():
                title_map[row[0]] = row[1]

        results = []
        seen = set()

        # Add content search results
        for sop_id, page, snippet, rank in content_rows:
            key = f"{sop_id}:{page}"
            if key not in seen:
                seen.add(key)
                results.append({
                    "sop_id": sop_id,
                    "page": page,
                    "snippet": snippet,
                    "rank": float(rank),
                    "title": title_map.get(sop_id, sop_id),
                })

        # Add sop-level matches that didn't appear in content results
        for sop_id, title in sop_rows:
            if not any(r["sop_id"] == sop_id for r in results):
                results.append({
                    "sop_id": sop_id,
                    "page": None,
                    "snippet": f"Document: <mark>{title or sop_id}</mark>",
                    "rank": 0.1,
                    "title": title or sop_id,
                })

        # Sort by rank descending
        results.sort(key=lambda r: r["rank"], reverse=True)
        return results[:limit]
    finally:
        conn.close()


@router.get("/stats")
async def get_stats(tenant_id: str):
    return db.get_stats(tenant_id=tenant_id)


@router.get("/starter-questions")
async def get_starter_questions(tenant_id: str, limit: int = Query(4)):
    """Get random starter questions from trained Q&A pairs and intent routes."""
    import random
    all_questions = []
    docs = db.list_sops(tenant_id=tenant_id)
    doc_titles = {d["sop_id"]: d.get("title", d["sop_id"]) for d in docs}

    # Source 1: Q&A pairs from documents
    for doc in docs:
        qa = doc.get("qa_pairs") or []
        if isinstance(qa, str):
            try:
                qa = json.loads(qa)
            except Exception:
                qa = []
        for pair in qa:
            q = pair.get("question", "") if isinstance(pair, dict) else ""
            if q and 15 < len(q) < 120:
                all_questions.append({"question": q, "sop_id": doc["sop_id"], "title": doc.get("title", "")})

    # Source 2: Intent routes (from training discovery + feedback)
    if len(all_questions) < limit * 2:
        conn = db.get_db(tenant_id)
        try:
            routes = conn.execute(
                "SELECT intent, sop_id FROM intent_routes WHERE source IN ('auto','discovered','feedback') ORDER BY hit_count DESC LIMIT 30"
            ).fetchall()
            for r in routes:
                q = r["intent"]
                if q and 15 < len(q) < 120 and not q.startswith("NEGATIVE:"):
                    all_questions.append({"question": q, "sop_id": r["sop_id"], "title": doc_titles.get(r["sop_id"], r["sop_id"])})
        finally:
            conn.close()

    if not all_questions:
        return []
    # Deduplicate by question text
    seen = set()
    unique = []
    for q in all_questions:
        key = q["question"].lower().strip()
        if key not in seen:
            seen.add(key)
            unique.append(q)
    random.shuffle(unique)
    return unique[:min(limit, 8)]


@router.get("/departments")
async def get_departments(tenant_id: str):
    return db.get_departments(tenant_id=tenant_id)


@router.get("/sops")
async def get_sops(tenant_id: str, department: Optional[str] = Query(None), search: Optional[str] = Query(None)):
    return db.list_sops(department=department, search=search, tenant_id=tenant_id)


# Static SOP routes MUST come before /sops/{sop_id} to avoid route conflicts
@router.get("/sops/compare")
async def compare_docs(tenant_id: str, old: str = Query(...), new: str = Query(...)):
    """Compare two standardized SOPs."""
    if not old or not new:
        return {"error": "Both 'old' and 'new' document IDs are required"}
    if old == new:
        return {"error": "Cannot compare a document with itself"}
    from backend.core.sop_compare import compare_sops
    return compare_sops(old, new, tenant_id=tenant_id)


@router.get("/sops/gap-scan")
async def gap_scan(tenant_id: str):
    """AI-powered library gap analysis."""
    import asyncio
    from backend.core.sop_gap_scanner import scan_library
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: scan_library(tenant_id=tenant_id))


@router.get("/sop-templates")
async def list_templates(tenant_id: str):
    """List available SOP templates."""
    from backend.core.sop_templates import get_templates
    return get_templates()


@router.get("/sop-templates/{template_id}")
async def get_template_detail(tenant_id: str, template_id: str):
    """Get full template definition."""
    from backend.core.sop_templates import get_template
    t = get_template(template_id)
    if not t:
        return {"error": "Template not found"}
    return t


@router.get("/sops/{sop_id}")
async def get_sop_detail(tenant_id: str, sop_id: str):
    sop = db.get_sop(sop_id, tenant_id=tenant_id)
    if not sop:
        return {"error": "Document not found"}
    page_contents = db.get_page_contents(sop_id, tenant_id=tenant_id)
    screenshots = db.get_screenshots(sop_id, tenant_id=tenant_id)
    compliance = db.get_compliance(sop_id, tenant_id=tenant_id)
    intent_routes = db.get_intent_routes(sop_id=sop_id, tenant_id=tenant_id)
    # Extract gap_analysis from standardized_json
    gap_analysis = None
    std_json = sop.get("standardized_json")
    if std_json:
        import json as _json
        if isinstance(std_json, str):
            try: std_json = _json.loads(std_json)
            except: std_json = None
        if isinstance(std_json, dict):
            gap_analysis = std_json.get("gap_analysis")
    return {**sop, "page_contents": page_contents, "compliance": compliance,
            "extracted_images": screenshots, "intent_routes": intent_routes,
            "gap_analysis": gap_analysis}


@router.get("/sops/{sop_id}/versions")
async def get_versions(tenant_id: str, sop_id: str):
    """Get version history by following the previous_version_id chain."""
    versions = []
    current_id = sop_id
    seen = set()
    while current_id and current_id not in seen:
        seen.add(current_id)
        conn = db.get_db(tenant_id)
        try:
            row = conn.execute(
                "SELECT sop_id, title, version, previous_version_id, created_at, page_count, department, is_enhanced FROM sops WHERE sop_id = %s",
                (current_id,)
            ).fetchone()
        finally:
            conn.close()
        if not row:
            break
        d = dict(row)
        d["created_at"] = d["created_at"].isoformat() if d.get("created_at") else None
        versions.append(d)
        current_id = d.get("previous_version_id")
    return {"sop_id": sop_id, "versions": versions, "total": len(versions)}


@router.put("/sops/{sop_id}/pin")
async def toggle_pin(tenant_id: str, sop_id: str):
    """Toggle pinned/starred state for a document."""
    conn = db.get_db(tenant_id)
    try:
        conn.execute("UPDATE sops SET pinned = NOT COALESCE(pinned, FALSE) WHERE sop_id = %s", (sop_id,))
        conn.commit()
        row = conn.execute("SELECT pinned FROM sops WHERE sop_id = %s", (sop_id,)).fetchone()
        return {"pinned": row["pinned"] if row else False}
    finally:
        conn.close()


@router.put("/sops/{sop_id}")
async def update_sop_fields(tenant_id: str, sop_id: str, body: dict = Body(default={})):
    """Update mutable SOP fields (currently: tags)."""
    sop = db.get_sop(sop_id, tenant_id=tenant_id)
    if not sop:
        return {"error": "Document not found"}
    conn = db.get_db(tenant_id)
    try:
        if "tags" in body:
            tags = body["tags"] if isinstance(body["tags"], list) else []
            conn.execute("UPDATE sops SET tags = %s WHERE sop_id = %s", (json.dumps(tags), sop_id))
            conn.commit()
        return {"status": "updated", "sop_id": sop_id}
    finally:
        conn.close()


@router.delete("/sops/{sop_id}")
async def delete_sop(tenant_id: str, sop_id: str):
    sop = db.get_sop(sop_id, tenant_id=tenant_id)
    if not sop:
        return {"error": "Document not found"}
    screenshot_dir = _tenant_data_dir(tenant_id) / "screenshots" / sop_id
    if screenshot_dir.exists():
        shutil.rmtree(screenshot_dir)
    db.delete_sop(sop_id, tenant_id=tenant_id)
    db.log_audit(tenant_id, "delete_document", resource_type="sop", resource_id=sop_id, details=sop.get("title", ""))
    return {"status": "deleted", "sop_id": sop_id}


@router.get("/sops/{sop_id}/pages/{page_num}")
async def get_page_image(tenant_id: str, sop_id: str, page_num: int, dpi: int = Query(200)):
    """Render a PDF page as PNG."""
    dpi = max(72, min(dpi, 600))  # Clamp to safe range
    import fitz
    sop = db.get_sop(sop_id, tenant_id=tenant_id)
    if not sop or not sop.get("pdf_path"):
        return {"error": "Document not found"}
    pdf_path = db.resolve_pdf_path(sop.get("pdf_path", ""))
    if not pdf_path:
        # Try tenant uploads dir
        pdf_path = str(_tenant_data_dir(tenant_id) / "uploads" / Path(sop.get("pdf_path", "")).name)
        if not Path(pdf_path).exists():
            return {"error": "PDF file not found"}
    doc = None
    try:
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        if page_num < 1 or page_num > total_pages:
            return Response(
                content=json.dumps({"error": f"Page {page_num} out of range", "total_pages": total_pages}),
                media_type="application/json", status_code=404
            )
        page = doc[page_num - 1]
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        return Response(content=img_bytes, media_type="image/png", headers={
            "Cache-Control": "public, max-age=86400",
            "X-Total-Pages": str(total_pages),
        })
    except Exception as e:
        return {"error": str(e)}
    finally:
        if doc:
            doc.close()


@router.get("/sops/{sop_id}/preview")
async def get_preview(tenant_id: str, sop_id: str):
    """Serve HTML preview — tries standardized DOCX first, then original file."""
    preview_path = _tenant_data_dir(tenant_id) / "previews" / f"{sop_id}.html"
    std_preview_path = _tenant_data_dir(tenant_id) / "previews" / f"{sop_id}_standardized.html"
    # Check if standardized preview exists or can be generated
    std_docx = _tenant_data_dir(tenant_id) / "standardized" / f"{sop_id}_standardized.docx"
    if std_docx.exists() and not std_preview_path.exists():
        std_preview_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            import mammoth
            with open(str(std_docx), "rb") as f:
                result = mammoth.convert_to_html(f)
                html = f"<!DOCTYPE html><html><head><meta charset='utf-8'><style>*{{border-radius:0!important}}body{{font-family:'Space Grotesk',sans-serif;padding:40px;max-width:900px;margin:0 auto;line-height:1.7;background:#feffd6;color:#383832}}h1,h2,h3{{font-weight:900;text-transform:uppercase;letter-spacing:-0.02em}}table{{border-collapse:collapse;width:100%;margin:16px 0}}th{{background:#383832;color:#feffd6;padding:8px 12px;text-transform:uppercase;font-size:11px;font-weight:900}}td{{border:2px solid #383832;padding:8px 12px}}tr:nth-child(even){{background:#fcf9ef}}strong{{font-weight:900}}ul{{padding-left:20px}}li{{margin:4px 0}}</style></head><body>{result.value}</body></html>"
                std_preview_path.write_text(html, encoding="utf-8")
        except Exception:
            pass
    if std_preview_path.exists():
        return Response(content=std_preview_path.read_text(encoding="utf-8"), media_type="text/html")
    # Fallback: try original file preview
    if not preview_path.exists():
        sop = db.get_sop(sop_id, tenant_id=tenant_id)
        if sop:
            file_path = sop.get("pdf_path", "")
            if not Path(file_path).exists():
                file_path = str(_tenant_data_dir(tenant_id) / "uploads" / Path(file_path).name)
            if Path(file_path).exists():
                ext = Path(file_path).suffix.lower()
                preview_path.parent.mkdir(parents=True, exist_ok=True)
                try:
                    if ext in (".docx", ".doc"):
                        import mammoth
                        with open(file_path, "rb") as f:
                            result = mammoth.convert_to_html(f)
                            html = f"<!DOCTYPE html><html><head><meta charset='utf-8'><style>body{{font-family:Inter,sans-serif;padding:40px;max-width:900px;margin:0 auto;line-height:1.7}}table{{border-collapse:collapse;width:100%;margin:16px 0}}th,td{{border:1px solid #e0e0e0;padding:8px 12px}}</style></head><body>{result.value}</body></html>"
                            preview_path.write_text(html, encoding="utf-8")
                    elif ext in (".xlsx", ".xls"):
                        from openpyxl import load_workbook
                        wb = load_workbook(file_path, data_only=True, read_only=True)
                        parts = ["<!DOCTYPE html><html><head><meta charset='utf-8'><style>body{font-family:Inter,sans-serif;padding:40px}table{border-collapse:collapse;width:100%;margin:16px 0}th,td{border:1px solid #e0e0e0;padding:8px 12px;font-size:13px}th{background:#f5f5f5}</style></head><body>"]
                        for sn in wb.sheetnames:
                            ws = wb[sn]
                            parts.append(f"<h2>{sn}</h2><table>")
                            for ri, row in enumerate(ws.iter_rows(values_only=True, max_row=500)):
                                cells = [str(c) if c else "" for c in row]
                                if not any(cells): continue
                                tag = "th" if ri == 0 else "td"
                                parts.append("<tr>" + "".join(f"<{tag}>{c}</{tag}>" for c in cells) + "</tr>")
                            parts.append("</table>")
                        parts.append("</body></html>")
                        preview_path.write_text("".join(parts), encoding="utf-8")
                        wb.close()
                except Exception:
                    pass
    if preview_path.exists():
        return Response(content=preview_path.read_text(encoding="utf-8"), media_type="text/html")
    return Response(
        content="<!DOCTYPE html><html><body style='font-family:sans-serif;padding:40px;text-align:center;color:#65655e;background:#feffd6'><p style='font-weight:700'>Preview not available for this document type.</p><p style='font-size:13px'>Use the Original PDF or Download options instead.</p></body></html>",
        media_type="text/html"
    )


@router.get("/sops/{sop_id}/download")
async def download_file(tenant_id: str, sop_id: str):
    """Download original uploaded file."""
    sop = db.get_sop(sop_id, tenant_id=tenant_id)
    if not sop:
        return {"error": "Document not found"}
    file_path = sop.get("pdf_path", "")
    if not Path(file_path).exists():
        file_path = str(_tenant_data_dir(tenant_id) / "uploads" / Path(file_path).name)
    if not Path(file_path).exists():
        return {"error": "File not found"}
    filename = Path(file_path).name.replace('"', '_')  # Sanitize for Content-Disposition
    with open(file_path, "rb") as f:
        content = f.read()
    return Response(content=content, media_type="application/octet-stream",
                    headers={"Content-Disposition": f'attachment; filename="{filename}"'})


@router.post("/upload")
async def upload_doc(tenant_id: str, file: UploadFile = File(...)):
    """Upload a document to tenant's storage."""
    if not file.filename:
        return {"error": "No filename"}
    allowed = (".pdf", ".docx", ".doc", ".xlsx", ".xls")
    ext = Path(file.filename).suffix.lower()
    if ext not in allowed:
        return {"error": f"Unsupported file type: {ext}"}

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        return {"error": "File exceeds 50MB limit"}

    safe_name = re.sub(r'[^\w\s\-.]', '', file.filename)
    if not safe_name:
        return {"error": "Invalid filename"}

    upload_dir = _tenant_data_dir(tenant_id) / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)

    # Check disk space before writing
    disk = shutil.disk_usage(upload_dir)
    if disk.free < len(content) + 100 * 1024 * 1024:  # 100MB buffer
        return {"error": f"Insufficient disk space. Available: {disk.free // (1024 * 1024)}MB"}

    save_path = upload_dir / safe_name

    try:
        with open(save_path, "wb") as f:
            f.write(content)
    except OSError as e:
        return {"error": f"Failed to save file: {e}"}

    sop_id = Path(safe_name).stem

    # ── Version handling: if sop_id already exists, archive the old version ──
    new_version = 1
    previous_version_id = None
    existing = db.get_sop(sop_id, tenant_id=tenant_id)
    if existing:
        old_version = existing.get("version") or 1
        new_version = old_version + 1
        versioned_id = f"{sop_id}_v{old_version}"
        previous_version_id = versioned_id
        # Rename old document's sop_id to versioned ID (preserves pages/embeddings/screenshots)
        conn = db.get_db(tenant_id)
        try:
            conn.execute("UPDATE sops SET sop_id = %s WHERE sop_id = %s", (versioned_id, sop_id))
            conn.execute("UPDATE page_content SET sop_id = %s WHERE sop_id = %s", (versioned_id, sop_id))
            conn.execute("UPDATE embeddings SET sop_id = %s WHERE sop_id = %s", (versioned_id, sop_id))
            conn.execute("UPDATE screenshots SET sop_id = %s WHERE sop_id = %s", (versioned_id, sop_id))
            conn.execute("UPDATE intent_routes SET sop_id = %s WHERE sop_id = %s", (versioned_id, sop_id))
            conn.execute("UPDATE compliance SET sop_id = %s WHERE sop_id = %s", (versioned_id, sop_id))
            conn.commit()
        finally:
            conn.close()

    # Get page count
    page_count = 0
    if ext == ".pdf":
        import PyPDF2
        try:
            reader = PyPDF2.PdfReader(str(save_path))
            page_count = len(reader.pages)
        except Exception:
            pass
    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(str(save_path))
            page_count = max(1, len([p for p in doc.paragraphs if p.text.strip()]) // 30)
        except Exception:
            page_count = 1
    elif ext in (".xlsx", ".xls"):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(save_path), read_only=True)
            page_count = len(wb.sheetnames)
            wb.close()
        except Exception:
            page_count = 1

    db.upsert_sop({
        "sop_id": sop_id, "title": file.filename, "description": "",
        "category_id": "", "department": "Uploaded", "system": "", "type": "",
        "tags": [], "pdf_path": str(save_path), "page_count": page_count,
        "doc_description": f"Uploaded file: {file.filename}",
        "total_screenshots": 0, "qa_pairs": [], "search_keywords": [],
        "entities": {}, "summary_short": "",
    }, tenant_id=tenant_id)

    # Set version metadata on the new document
    conn = db.get_db(tenant_id)
    try:
        conn.execute(
            "UPDATE sops SET version = %s, previous_version_id = %s WHERE sop_id = %s",
            (new_version, previous_version_id, sop_id)
        )
        conn.commit()
    finally:
        conn.close()

    db.log_audit(tenant_id, "upload_document", resource_type="sop", resource_id=sop_id,
                 details=f"{file.filename} ({page_count}p) v{new_version}")
    return {"sop_id": sop_id, "filename": file.filename, "pages": page_count,
            "version": new_version, "status": "uploaded"}


@router.post("/upload-multiple")
async def upload_multiple(tenant_id: str, files: list[UploadFile] = File(...)):
    """Upload multiple documents to tenant's storage in one request."""
    if not files:
        return {"error": "No files provided", "results": []}

    allowed = (".pdf", ".docx", ".doc", ".xlsx", ".xls")
    max_size = 50 * 1024 * 1024  # 50MB per file
    results: list[dict] = []

    # Read all file contents upfront for disk space check
    file_data: list[tuple[UploadFile, bytes, str, str]] = []  # (file, content, ext, safe_name)
    total_size = 0
    for f in files:
        if not f.filename:
            results.append({"sop_id": None, "filename": "(no filename)", "status": "error", "error": "No filename"})
            continue
        ext = Path(f.filename).suffix.lower()
        if ext not in allowed:
            results.append({"sop_id": None, "filename": f.filename, "status": "error", "error": f"Unsupported file type: {ext}"})
            continue
        content = await f.read()
        if len(content) > max_size:
            results.append({"sop_id": None, "filename": f.filename, "status": "error", "error": "File exceeds 50MB limit"})
            continue
        safe_name = re.sub(r'[^\w\s\-.]', '', f.filename)
        if not safe_name:
            results.append({"sop_id": None, "filename": f.filename, "status": "error", "error": "Invalid filename"})
            continue
        total_size += len(content)
        file_data.append((f, content, ext, safe_name))

    if not file_data:
        return {"results": results}

    # Check disk space once for total size
    upload_dir = _tenant_data_dir(tenant_id) / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)
    disk = shutil.disk_usage(upload_dir)
    if disk.free < total_size + 100 * 1024 * 1024:  # 100MB buffer
        return {"error": f"Insufficient disk space. Need {total_size // (1024*1024)}MB, available: {disk.free // (1024*1024)}MB", "results": results}

    # Save each file and create DB records
    for f, content, ext, safe_name in file_data:
        save_path = upload_dir / safe_name
        try:
            with open(save_path, "wb") as fh:
                fh.write(content)
        except OSError as e:
            results.append({"sop_id": None, "filename": f.filename, "status": "error", "error": f"Failed to save: {e}"})
            continue

        sop_id = Path(safe_name).stem
        page_count = 0
        if ext == ".pdf":
            import PyPDF2
            try:
                reader = PyPDF2.PdfReader(str(save_path))
                page_count = len(reader.pages)
            except Exception:
                pass
        elif ext in (".docx", ".doc"):
            try:
                from docx import Document
                doc = Document(str(save_path))
                page_count = max(1, len([p for p in doc.paragraphs if p.text.strip()]) // 30)
            except Exception:
                page_count = 1
        elif ext in (".xlsx", ".xls"):
            try:
                from openpyxl import load_workbook
                wb = load_workbook(str(save_path), read_only=True)
                page_count = len(wb.sheetnames)
                wb.close()
            except Exception:
                page_count = 1

        db.upsert_sop({
            "sop_id": sop_id, "title": f.filename, "description": "",
            "category_id": "", "department": "Uploaded", "system": "", "type": "",
            "tags": [], "pdf_path": str(save_path), "page_count": page_count,
            "doc_description": f"Uploaded file: {f.filename}",
            "total_screenshots": 0, "qa_pairs": [], "search_keywords": [],
            "entities": {}, "summary_short": "",
        }, tenant_id=tenant_id)

        db.log_audit(tenant_id, "upload_document", resource_type="sop", resource_id=sop_id, details=f"{f.filename} ({page_count}p)")
        results.append({"sop_id": sop_id, "filename": f.filename, "pages": page_count, "status": "uploaded"})

    return {"results": results}


@router.post("/process/{sop_id}")
async def process_doc(tenant_id: str, sop_id: str):
    """Run full pipeline + auto-train for a tenant document."""
    import asyncio
    from backend.core.trainer import process_and_train

    # Prevent concurrent processing of the same document
    doc_key = f"{tenant_id}:{sop_id}"
    with _processing_lock:
        if doc_key in _processing_docs:
            return {"error": "Document is already being processed. Please wait."}
        _processing_docs.add(doc_key)

    try:
        sop = db.get_sop(sop_id, tenant_id=tenant_id)
        if not sop:
            return {"error": "Document not found"}

        pdf_path = sop.get("pdf_path", "")
        if not Path(pdf_path).exists():
            pdf_path = str(_tenant_data_dir(tenant_id) / "uploads" / Path(pdf_path).name)
        if not Path(pdf_path).exists():
            return {"error": "File not found"}

        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, lambda: process_and_train(pdf_path, sop_id, tenant_id=tenant_id))
        db.log_audit(tenant_id, "process_document", resource_type="sop", resource_id=sop_id, details=f"{sop.get('title', '')} — {result.get('page_count', 0)} pages")
        return result
    finally:
        with _processing_lock:
            _processing_docs.discard(doc_key)


@router.post("/process/{sop_id}/stop")
async def stop_processing(tenant_id: str, sop_id: str):
    """Stop a running pipeline for a document."""
    doc_key = f"{tenant_id}:{sop_id}"
    with _processing_lock:
        _processing_docs.discard(doc_key)
    # Signal the trainer to stop
    from backend.core.trainer import stop_training
    stop_training()
    return {"status": "stopped", "sop_id": sop_id}


@router.post("/process/{sop_id}/stream")
async def process_doc_stream(tenant_id: str, sop_id: str):
    """Run full pipeline + auto-train with real-time SSE status streaming."""
    from backend.core.trainer import process_and_train

    # Prevent concurrent processing of the same document
    doc_key = f"{tenant_id}:{sop_id}"
    with _processing_lock:
        if doc_key in _processing_docs:
            return {"error": "Document is already being processed. Please wait."}
        _processing_docs.add(doc_key)

    sop = db.get_sop(sop_id, tenant_id=tenant_id)
    if not sop:
        with _processing_lock:
            _processing_docs.discard(doc_key)
        return {"error": "Document not found"}

    pdf_path = sop.get("pdf_path", "")
    if not Path(pdf_path).exists():
        pdf_path = str(_tenant_data_dir(tenant_id) / "uploads" / Path(pdf_path).name)
    if not Path(pdf_path).exists():
        with _processing_lock:
            _processing_docs.discard(doc_key)
        return {"error": "File not found"}

    status_queue: Queue = Queue()

    def on_status(step: str, msg: str):
        status_queue.put_nowait({"step": step, "message": msg})

    def run_pipeline():
        return process_and_train(pdf_path, sop_id, on_status=on_status, tenant_id=tenant_id)

    async def event_generator():
        loop = asyncio.get_event_loop()
        # Run pipeline in background — do NOT cancel on client disconnect
        task = asyncio.ensure_future(loop.run_in_executor(None, run_pipeline))

        # Background cleanup: when task finishes, always release the lock
        def _on_task_done(t):
            with _processing_lock:
                _processing_docs.discard(doc_key)
            try:
                result = t.result()
                db.log_audit(tenant_id, "process_document", resource_type="sop", resource_id=sop_id,
                             details=f"{sop.get('title', '')} — {result.get('page_count', 0)} pages")
            except Exception:
                pass
        task.add_done_callback(_on_task_done)

        try:
            while not task.done():
                try:
                    await asyncio.sleep(0.3)
                    while not status_queue.empty():
                        status = status_queue.get_nowait()
                        yield f"event: status\ndata: {json.dumps(status)}\n\n"
                except (asyncio.CancelledError, GeneratorExit):
                    # Client disconnected — pipeline keeps running in background
                    return

            # Drain remaining status events
            while not status_queue.empty():
                status = status_queue.get_nowait()
                yield f"event: status\ndata: {json.dumps(status)}\n\n"

            result = task.result()
            yield f"event: done\ndata: {json.dumps({'result': result})}\n\n"

        except Exception as e:
            yield f"event: error\ndata: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@router.get("/training/logs")
async def get_training_logs(tenant_id: str, since: int = 0):
    """Get real-time training logs."""
    from backend.core.trainer import get_training_logs, get_training_status
    return {"logs": get_training_logs(since), "status": get_training_status(), "total": len(get_training_logs())}


@router.get("/analytics")
async def get_analytics(tenant_id: str):
    """Analytics dashboard data: daily counts, top queries, low quality queries."""
    conn = db.get_db(tenant_id)
    try:
        # Daily query counts (last 7 days)
        daily = conn.execute("""
            SELECT DATE(created_at) as day, COUNT(*) as count,
                   AVG(CASE WHEN duration_s IS NOT NULL THEN duration_s ELSE NULL END) as avg_duration
            FROM query_log
            WHERE created_at > NOW() - INTERVAL '7 days'
            GROUP BY DATE(created_at) ORDER BY day
        """).fetchall()
        # Top queries (grouped by question text)
        top = conn.execute("""
            SELECT question, COUNT(*) as count
            FROM query_log GROUP BY question ORDER BY count DESC LIMIT 10
        """).fetchall()
        # Low quality (low score or thumbs down)
        low = conn.execute("""
            SELECT question, quality_score, feedback, created_at
            FROM query_log WHERE quality_score < 50 OR feedback = 'down'
            ORDER BY created_at DESC LIMIT 10
        """).fetchall()
        # Satisfaction stats
        thumbs_up = conn.execute("SELECT COUNT(*) FROM query_log WHERE feedback = 'up'").fetchone()[0]
        thumbs_down = conn.execute("SELECT COUNT(*) FROM query_log WHERE feedback = 'down'").fetchone()[0]
        total_queries = conn.execute("SELECT COUNT(*) FROM query_log").fetchone()[0]
        low_quality_count = conn.execute("SELECT COUNT(*) FROM query_log WHERE quality_score < 50").fetchone()[0]
        return {
            "daily": [{"day": str(r[0]), "count": r[1], "avg_duration": float(r[2]) if r[2] else 0} for r in daily],
            "top_queries": [{"question": r[0], "count": r[1]} for r in top],
            "low_quality": [{"question": r[0], "quality_score": r[1], "feedback": r[2], "created_at": str(r[3])} for r in low],
            "thumbs_up": thumbs_up or 0,
            "thumbs_down": thumbs_down or 0,
            "total_queries": total_queries or 0,
            "low_quality_count": low_quality_count or 0,
        }
    finally:
        conn.close()


@router.get("/logs")
async def get_logs(tenant_id: str, limit: int = Query(50)):
    return db.get_query_logs(limit, tenant_id=tenant_id)


@router.get("/logs/downvoted")
async def get_downvoted(tenant_id: str, limit: int = Query(20)):
    """Get queries with negative feedback for admin review."""
    logs = db.get_query_logs(200, tenant_id=tenant_id)
    return [l for l in logs if l.get("feedback") == "down"][:limit]


# ── SOP Standardization ─────────────────────────────────────────────────────

@router.post("/sops/{sop_id}/standardize")
async def standardize_doc(tenant_id: str, sop_id: str):
    """Run document standardization on a processed document."""
    import asyncio
    from backend.core.sop_standardize import standardize_sop
    sop = db.get_sop(sop_id, tenant_id=tenant_id)
    if not sop:
        return {"error": "Document not found"}
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, lambda: standardize_sop(sop_id, tenant_id=tenant_id))
    db.log_audit(tenant_id, "standardize_document", resource_type="sop", resource_id=sop_id, details=sop.get("title", ""))
    return result


@router.get("/sops/{sop_id}/standardized")
async def get_standardized(tenant_id: str, sop_id: str):
    """Get standardized SOP metadata."""
    sop = db.get_sop(sop_id, tenant_id=tenant_id)
    if not sop:
        return {"error": "Document not found"}
    return {
        "sop_id": sop_id,
        "sop_score": sop.get("sop_score", 0),
        "sop_gaps": sop.get("sop_gaps"),
        "has_standardized": sop.get("standardized_json") is not None,
    }


@router.get("/sops/{sop_id}/download/docx")
async def download_docx(tenant_id: str, sop_id: str):
    """Download standardized document as DOCX."""
    docx_path = _tenant_data_dir(tenant_id) / "standardized" / f"{sop_id}_standardized.docx"
    if not docx_path.exists():
        return {"error": "Standardized DOCX not found. Run standardize first."}
    return Response(
        content=docx_path.read_bytes(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{sop_id}_standardized.docx"'}
    )


@router.get("/sops/{sop_id}/download/pdf")
async def download_pdf(tenant_id: str, sop_id: str):
    """Download standardized SOP as PDF (converted from DOCX via LibreOffice)."""
    docx_path = _tenant_data_dir(tenant_id) / "standardized" / f"{sop_id}_standardized.docx"
    pdf_path = _tenant_data_dir(tenant_id) / "standardized" / f"{sop_id}_standardized.pdf"

    if not docx_path.exists():
        return {"error": "Standardized DOCX not found. Run standardize first."}

    # Convert DOCX → PDF via LibreOffice headless (if available)
    if not pdf_path.exists():
        import subprocess
        try:
            out_dir = str(docx_path.parent)
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, str(docx_path)],
                capture_output=True, timeout=60
            )
        except FileNotFoundError:
            return {"error": "PDF conversion requires LibreOffice. Download the DOCX instead."}
        except Exception as e:
            return {"error": f"PDF conversion failed: {e}"}

    if pdf_path.exists():
        return Response(
            content=pdf_path.read_bytes(),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{sop_id}_standardized.pdf"'}
        )
    return {"error": "PDF conversion failed"}



# (compare and gap-scan routes moved above /sops/{sop_id} to avoid route conflicts)


# ── Multi-Language ───────────────────────────────────────────────────────────

@router.post("/sops/{sop_id}/translate/{lang}")
async def translate_doc(tenant_id: str, sop_id: str, lang: str):
    """Translate standardized document to another language. lang = 'my', 'zh', 'ja', etc."""
    import asyncio
    from backend.core.sop_translate import translate_sop
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: translate_sop(sop_id, lang, tenant_id=tenant_id))


@router.get("/sops/{sop_id}/translations")
async def list_translations(tenant_id: str, sop_id: str):
    """List available translations for a document."""
    from backend.core.sop_translate import get_available_translations, SUPPORTED_LANGUAGES
    return {
        "available": get_available_translations(sop_id, tenant_id=tenant_id),
        "supported_languages": SUPPORTED_LANGUAGES,
    }


@router.get("/sops/{sop_id}/download/{lang}/docx")
async def download_translated_docx(tenant_id: str, sop_id: str, lang: str):
    """Download translated DOCX."""
    if lang == "en":
        docx_path = _tenant_data_dir(tenant_id) / "standardized" / f"{sop_id}_standardized.docx"
    else:
        docx_path = _tenant_data_dir(tenant_id) / "standardized" / f"{sop_id}_standardized_{lang}.docx"
    if not docx_path.exists():
        return {"error": f"Translation not found for language: {lang}. Generate it first."}
    return Response(
        content=docx_path.read_bytes(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{sop_id}_{lang}.docx"'}
    )



# ── Categories ───────────────────────────────────────────────────────────────

@router.get("/categories")
async def get_categories(tenant_id: str):
    return db.get_categories(tenant_id=tenant_id)


# ── Evals ────────────────────────────────────────────────────────────────────

@router.post("/run-evals")
async def run_evals(tenant_id: str, category: Optional[str] = Query(None)):
    import asyncio
    from backend.evals.run_evals import run_all_evals
    loop = asyncio.get_event_loop()
    results = await loop.run_in_executor(None, lambda: run_all_evals(category, tenant_id=tenant_id))
    total = results.get("total", 0)
    passed = results.get("passed", 0)
    failed = total - passed
    score = round((passed / total * 100) if total > 0 else 0, 1)
    db.save_eval_run(category or "all", total, passed, failed, score, results.get("results", []), tenant_id=tenant_id)
    return results


@router.get("/evals/history")
async def get_eval_history(tenant_id: str):
    return db.get_eval_history(tenant_id=tenant_id)


@router.get("/evals/test-cases")
async def get_eval_test_cases(tenant_id: str):
    from backend.evals.test_cases import get_test_cases
    cases = get_test_cases()
    return [{"question": tc.question, "expected_strings": tc.expected_strings,
             "category": tc.category, "golden_doc": tc.golden_doc,
             "description": tc.description} for tc in cases]


# ── Knowledge Extraction ────────────────────────────────────────────────────

@router.post("/extract-knowledge")
async def extract_knowledge_all(tenant_id: str):
    import asyncio
    from backend.core.knowledge_extract import extract_all_knowledge
    loop = asyncio.get_event_loop()
    results = await loop.run_in_executor(None, lambda: extract_all_knowledge(tenant_id=tenant_id))
    return {"extracted": len(results), "results": results}


@router.post("/extract-knowledge/{sop_id}")
async def extract_knowledge_single(tenant_id: str, sop_id: str):
    import asyncio
    from backend.core.knowledge_extract import extract_knowledge
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, lambda: extract_knowledge(sop_id, tenant_id=tenant_id))
    return result


# ── Relationships ────────────────────────────────────────────────────────────

@router.get("/relationships/{sop_id}")
async def get_related_docs(tenant_id: str, sop_id: str):
    relationships = db.get_relationships(sop_id, tenant_id=tenant_id)
    related = db.find_related_documents(sop_id, tenant_id=tenant_id)
    return {"relationships": relationships, "related": related}


# ── Compliance ───────────────────────────────────────────────────────────────

@router.get("/compliance")
async def get_compliance_all(tenant_id: str):
    return db.get_all_compliance(tenant_id=tenant_id)


@router.post("/compliance/check-all")
async def check_all_compliance(tenant_id: str):
    import asyncio
    from backend.core.compliance import check_all_compliance as _check_all
    loop = asyncio.get_event_loop()
    results = await loop.run_in_executor(None, lambda: _check_all(tenant_id=tenant_id))
    return {"checked": len(results), "results": results}


@router.post("/compliance/{sop_id}")
async def check_compliance_single(tenant_id: str, sop_id: str):
    import asyncio
    from backend.core.compliance import check_compliance
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, lambda: check_compliance(sop_id, tenant_id=tenant_id))
    return result


# ── Wiki Knowledge Layer ──────────────────────────────────────────────

@router.get("/wiki")
async def list_wiki(tenant_id: str):
    """List all wiki pages for this tenant."""
    return db.list_wiki_pages(tenant_id=tenant_id)


@router.get("/wiki/{page_id}")
async def get_wiki_page(tenant_id: str, page_id: str):
    """Get full wiki page content."""
    page = db.get_wiki_page(page_id, tenant_id=tenant_id)
    if not page:
        return {"error": "Wiki page not found"}
    return page


@router.post("/wiki/lint")
async def lint_wiki(tenant_id: str):
    """Run wiki health check — stale sources, orphans, contradictions."""
    from backend.core.wiki import wiki_lint
    return wiki_lint(tenant_id=tenant_id)


@router.delete("/wiki/{page_id}")
async def delete_wiki(tenant_id: str, page_id: str):
    """Delete a wiki page. Agent will recreate it on next document process if needed."""
    db.delete_wiki_page(page_id, tenant_id=tenant_id)
    return {"status": "deleted", "page_id": page_id}


@router.get("/schedule")
async def get_schedule(tenant_id: str):
    """Get scheduled re-training configuration."""
    retrain_enabled = db.get_runtime_config("retrain_enabled", tenant_id=tenant_id)
    retrain_interval = db.get_runtime_config("retrain_interval_days", tenant_id=tenant_id)
    last_retrain = db.get_runtime_config("last_retrain", tenant_id=tenant_id)
    return {
        "retrain_enabled": retrain_enabled == "true" if retrain_enabled else False,
        "retrain_interval_days": int(retrain_interval) if retrain_interval else 7,
        "last_retrain": last_retrain if last_retrain else None,
    }


@router.put("/schedule")
async def set_schedule(tenant_id: str, request: dict):
    """Set scheduled re-training configuration."""
    enabled = request.get("enabled", False)
    interval = request.get("interval_days", 7)
    db.set_runtime_config("retrain_enabled", str(enabled).lower(), tenant_id=tenant_id)
    db.set_runtime_config("retrain_interval_days", str(interval), tenant_id=tenant_id)
    return {"status": "saved", "retrain_enabled": enabled, "retrain_interval_days": interval}


@router.post("/generate-persona")
async def gen_persona(tenant_id: str):
    """Auto-generate agent persona from all documents."""
    import asyncio
    from backend.core.wiki import generate_persona
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: generate_persona(tenant_id))
